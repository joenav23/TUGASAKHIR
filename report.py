from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import date
import subprocess

def add_header(document):
    # Add Header
    header_section = document.sections[0]
    header = header_section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Add the logo
    logo_path = "Logo Horizontal.png"
    logo_height = Cm(1.5)
    run = header_paragraph.add_run()
    run.add_picture(logo_path, height=logo_height)
    
    # Add space after header
    header.add_paragraph()

def create_cover_page(document):
    add_header(document)
    
    # Add the logo image
    document.add_picture('tri.png', width=Inches(2))
    document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center align the image
    
    # Add spacer
    document.add_paragraph()
    document.paragraphs[-1].add_run().add_break()
    
    while len(document.paragraphs) % 4 != 0:
        document.add_paragraph()
    
    title = """Sample Penetration Test Report
    Example Company"""
    document.add_heading(title, level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center align the text
    
    # Add spacer
    document.add_paragraph()
    document.paragraphs[-1].add_run().add_break()
    
    while len(document.paragraphs) % 14 != 0:
        document.add_paragraph()
    
    # Add details
    text = f"""
    Company: Customer Name
    Date: {date.today().strftime('%d %B %Y')}
    Version 1.0"""
    document.add_paragraph(text)
    



def generate_findings(document):
    # Add the introduction section
    document.add_heading("Introduction", level=2)
    intro_paragraph = document.add_paragraph()
    intro_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    intro_text = """Laporan ini disusun sebagai hasil pengujian penetrasi terhadap CVE-2022-46169, sebuah kerentanan yang ditemukan dalam perangkat lunak Cacti. Kerentanan ini memungkinkan serangan tanpa autentikasi untuk melakukan eksekusi perintah sistem secara sewenang-wenang pada server yang menjalankan Cacti. Dalam laporan ini, kami akan memberikan detail mengenai temuan kerentanan CVE-2022-46169 yang kami identifikasi selama penilaian. Kami akan menjelaskan dengan rinci potensi dampak dan risiko yang terkait dengan kerentanan ini, serta memberikan rekomendasi tindakan untuk memperbaiki kerentanan tersebut dan meningkatkan keamanan sistem secara menyeluruh."""
    intro_paragraph.add_run(intro_text)
    
    # Add the executive summary section
    vuln_devices = []
    scan_devices = []
    with open("nmap_results.txt", 'r') as f_vuln:
        lines = f_vuln.readlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if "Nmap scan report for" in line:
            ip_address = line.split()[-1]
            cacti_detected = False
            for j in range(i, min(i + 10, len(lines))):  # Periksa 10 baris berikutnya untuk keberadaan Cacti
                if "http-title: Login to Cacti" in lines[j]:
                    cacti_detected = True
                    break
            if cacti_detected:
                scan_devices.append((ip_address, "Cacti Detected"))
            else:
                scan_devices.append((ip_address, "No Cacti Detected"))
        i += 1

    with open("vuln_scan.txt", 'r') as f_vuln:
        lines = f_vuln.readlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if "========== Vulnerability Scan Result for" in line:
            ip_address = line.split()[-1]
            if i + 5 < len(lines):
                status_line = lines[i+5].strip()
                if "The target appears to be vulnerable. The target is Cacti version 1.2.22" in status_line:
                    vuln_devices.append((ip_address, "CVE-2022-46169 Vulnerability Detected"))
                else:
                    vuln_devices.append((ip_address, "CVE-2022-46169 Vulnerability Detected"))
        i += 1

    document.add_heading("Executive Summary", level=2)
    if scan_devices:
        document.add_paragraph("Scanned Devices:")
    for device in scan_devices:
        document.add_paragraph(f"- {device[0]}: {device[1]}")
    if vuln_devices:
        document.add_paragraph("Vulnerable Devices:")
    for device in vuln_devices:
        document.add_paragraph(f"- {device[0]}: {device[1]}")

    summary_paragraph = document.add_paragraph()
    summary_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    summary_text = "Pada " + date.today().strftime('%d %B %Y') + ", kami melaksanakan uji penetrasi dengan menggunakan pengetahuan sebelumnya tentang lingkungan internal atau dengan menggunakan kredensial yang kami miliki sebelumnya. Tujuannya adalah untuk menemukan kelemahan keamanan dan menguji kemungkinan pengeksploitasian celah tersebut. Pengujian dilakukan secara otomatis dengan menggunakan alat khusus yang dirancang untuk mendeteksi kerentanan CVE-2022-46169 pada Cacti. Kerentanan CVE-2022-46169 merupakan kerentanan eksekusi kode dari jarak jauh pada Cacti. Kerentanan ini terjadi karena adanya kecacatan pada file remote_agent.php. File ini dapat diakses tanpa perlu otentikasi, sehingga dapat dimanfaatkan oleh penyerang untuk melakukan eksekusi kode dari jarak jauh."""
    summary_paragraph.add_run(summary_text)


    # Add the Nmap scan section
    document.add_heading("Scan Target", level=2)
    with open("nmap_results.txt", "r") as scan_file:
        scan_content = scan_file.read()
    document.add_paragraph(scan_content)

    # Add the Proof of Concept section
    document.add_heading("Exploitabel", level=2)
    with open("exploit.txt", "r") as exploit_file:
        exploit_content = exploit_file.read()
    document.add_paragraph(exploit_content)

    # Add the Recommendations section
    document.add_heading("Recommendation", level=2)
    recommendations_paragraph = document.add_paragraph()
    recommendations_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    recommendations_text = "Untuk mengurangi risiko dari CVE-2022-46169, disarankan untuk mengambil langkah-langkah berikut:"
    recommendations_paragraph.add_run(recommendations_text)


    # Add the recommendations list
    recommendations_list = [
        "Memperbarui Cacti ke versi terbaru yang tersedia.",
        "Menerapkan aturan firewall yang membatasi akses ke layanan Cacti.",
        "Melakukan evaluasi keamanan secara berkala dan pengujian penetrasi untuk mengidentifikasi dan mengatasi kerentanan.",
        "Menerapkan kebijakan sandi yang kuat dan menghindari penggunaan kredensial default.",
        "Rutin memperbarui perangkat untuk mengatasi masalah keamanan."
    ]
    for recommendation in recommendations_list:
        document.add_paragraph(recommendation, style='List Bullet')
    
    # Save the document
    document.save("Penetration_Test_Report.docx")

def generate_report():
    # Create a new Word document
    document = Document()

    # Generate the cover page
    create_cover_page(document)

    # Generate the findings page
    generate_findings(document)

    subprocess.run(["unoconv", "-f", "pdf", "Penetration_Test_Report.docx"], stderr=subprocess.DEVNULL)

    print("Penetration test report generated successfully.")

if __name__ == "__main__":
    generate_report()
